import re
import os
from docx import Document

# Function to extract marks from question text. Expected format: ... (Marks: X)
def get_marks(question_text):
    match = re.search(r'\(Marks\s*:\s*(\d+)\)', question_text)
    return int(match.group(1)) if match else 0

def list_files(directory, extension='.docx'):
    """List all files with the given extension in the directory"""
    files = [f for f in os.listdir(directory) if f.endswith(extension)]
    return files

def select_file(prompt, directory='.', extension='.docx'):
    """Prompt user to select a file from the current directory"""
    files = list_files(directory, extension)
    
    if not files:
        print(f"No {extension} files found in the current directory.")
        return None
    
    print(f"\n{prompt}")
    for idx, filename in enumerate(files, 1):
        print(f"{idx}. {filename}")
    
    while True:
        try:
            choice = int(input("\nEnter the number of your choice: "))
            if 1 <= choice <= len(files):
                return os.path.join(directory, files[choice-1])
            else:
                print("Invalid choice. Please try again.")
        except ValueError:
            print("Please enter a number.")

# Get the question bank file
question_bank_path = select_file("Select a question bank file:")
if not question_bank_path:
    print("No question bank file selected. Exiting.")
    exit()

# Load the question bank document
question_bank_doc = Document(question_bank_path)
questions = [para.text.strip() for para in question_bank_doc.paragraphs if para.text.strip()]

# Display questions for selection
print("\nQuestion Bank:")
for idx, q in enumerate(questions, start=1):
    print(f"{idx}. {q}")

# Let the user choose questions for each part
selected_A = input("\nEnter question numbers for PART A (comma separated): ")
selected_B = input("Enter question numbers for PART B (comma separated): ")

# Convert input to indices (0-indexed)
selected_A_idx = [int(x.strip()) - 1 for x in selected_A.split(',') if x.strip().isdigit()]
selected_B_idx = [int(x.strip()) - 1 for x in selected_B.split(',') if x.strip().isdigit()]

# Build lists of (question text, marks) for each part
selected_A_questions = [(questions[i], get_marks(questions[i])) for i in selected_A_idx]
selected_B_questions = [(questions[i], get_marks(questions[i])) for i in selected_B_idx]

# Get the template file
template_path = select_file("Select a template file:")
if not template_path:
    print("No template file selected. Exiting.")
    exit()

# Load the template DOCX file
template_doc = Document(template_path)
table = template_doc.tables[0]

# Remove existing rows except the header row (assumes the first row is the header)
while len(table.rows) > 1:
    table._tbl.remove(table.rows[1]._tr)

# Function to add a merged header row to the table
def add_section_header(table, header_text):
    row = table.add_row()
    merged_cell = row.cells[0]
    # Merge all cells in the row to create a section header
    for cell in row.cells[1:]:
        merged_cell = merged_cell.merge(cell)
    merged_cell.text = header_text

# Add PART A section header
add_section_header(table, "PART A")

# Add rows for PART A questions
for idx, (q_text, q_marks) in enumerate(selected_A_questions, start=1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)         # Q.No for PART A
    row_cells[1].text = q_text            # Question text
    row_cells[2].text = str(q_marks)      # Marks

# Add PART B section header
add_section_header(table, "PART B")

# Add rows for PART B questions
for idx, (q_text, q_marks) in enumerate(selected_B_questions, start=1):
    row_cells = table.add_row().cells
    row_cells[0].text = str(idx)         # Q.No for PART B
    row_cells[1].text = q_text            # Question text
    row_cells[2].text = str(q_marks)      # Marks

# Get output filename
output_filename = input("\nEnter the output filename (e.g., 'generated_question_paper.docx'): ")
if not output_filename.endswith('.docx'):
    output_filename += '.docx'

# Save the generated question paper
template_doc.save(output_filename)
print(f"\nQuestion paper saved as: {os.path.abspath(output_filename)}")