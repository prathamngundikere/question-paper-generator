import re
import os
import sys
from docx import Document

# Function to extract information from question text
def extract_question_info(question_text):
    # Extract marks
    marks_match = re.search(r'\(Marks\s*:\s*(\d+)\)', question_text)
    marks = int(marks_match.group(1)) if marks_match else 0
    
    # Extract CO
    co_match = re.search(r'\(CO\s*:\s*([^)]+)\)', question_text)
    co = co_match.group(1) if co_match else ""
    
    # Extract RBT
    rbt_match = re.search(r'\(RBT\s*:\s*([^)]+)\)', question_text)
    rbt = rbt_match.group(1) if rbt_match else ""
    
    # Clean question text by removing the marks, CO, and RBT information
    clean_question = question_text
    if marks_match:
        clean_question = clean_question.replace(marks_match.group(0), "")
    if co_match:
        clean_question = clean_question.replace(co_match.group(0), "")
    if rbt_match:
        clean_question = clean_question.replace(rbt_match.group(0), "")
    
    # Trim any extra spaces
    clean_question = clean_question.strip()
    
    return {
        "question": clean_question,
        "marks": marks,
        "co": co,
        "rbt": rbt
    }

def list_files(directory, extension='.docx'):
    """List all files with the given extension in the directory"""
    files = [f for f in os.listdir(directory) if f.endswith(extension)]
    return files

def select_file(prompt, directory='.', extension='.docx'):
    """Prompt user to select a file from the current directory"""
    files = list_files(directory, extension)
    
    if not files:
        print(f"No {extension} files found in the current directory.")
        return None
    
    print(f"\n{prompt}")
    for idx, filename in enumerate(files, 1):
        print(f"{idx}. {filename}")
    
    while True:
        try:
            choice = int(input("\nEnter the number of your choice: "))
            if 1 <= choice <= len(files):
                return os.path.join(directory, files[choice-1])
            else:
                print("Invalid choice. Please try again.")
        except ValueError:
            print("Please enter a number.")

def assign_question_parts(questions, target_marks=25):
    """
    Group questions into parts (a, b, c, etc.) aiming for a target total marks
    Returns a list of question groups, each with total marks close to target_marks
    """
    question_groups = []
    current_group = []
    current_marks = 0
    
    for question in questions:
        if current_marks + question["marks"] > target_marks and current_group:
            # If adding this question exceeds the target and we already have questions,
            # save the current group and start a new one
            question_groups.append(current_group)
            current_group = [question]
            current_marks = question["marks"]
        else:
            # Add the question to the current group
            current_group.append(question)
            current_marks += question["marks"]
    
    # Add any remaining questions
    if current_group:
        question_groups.append(current_group)
    
    return question_groups

def main():
    print("=" * 50)
    print("QUESTION PAPER GENERATOR")
    print("=" * 50)
    print("\nNOTE: This will create a question paper with:")
    print("- PART A: Two sets of questions (Q1 & Q2) worth 25 marks each")
    print("- PART B: Two sets of questions (Q3 & Q4) worth 25 marks each")
    print("- Students choose between Q1 or Q2, and between Q3 or Q4")
    print("=" * 50)
    
    # Get the question bank file
    question_bank_path = select_file("Select a question bank file:")
    if not question_bank_path:
        print("No question bank file selected. Exiting.")
        return
    
    # Load the question bank document
    try:
        question_bank_doc = Document(question_bank_path)
        raw_questions = [para.text.strip() for para in question_bank_doc.paragraphs if para.text.strip()]
        
        # Process questions to extract marks, CO, and RBT
        all_questions = [extract_question_info(q) for q in raw_questions]
        
        # Filter out questions with no marks
        all_questions = [q for q in all_questions if q["marks"] > 0]
        
    except Exception as e:
        print(f"Error loading question bank: {e}")
        input("Press Enter to exit...")
        return
    
    # Display questions with marks
    print("\nQuestion Bank:")
    for idx, q in enumerate(all_questions, start=1):
        print(f"{idx}. {q['question']} [{q['marks']} marks]")
    
    # Let the user choose questions for PART A (Q1)
    print("\n" + "=" * 50)
    print("PART A - Question 1 Selection")
    print("=" * 50)
    selected_A1 = input("Enter question numbers for Q1 (comma separated, total ~25 marks): ")
    
    # Convert input to indices (0-indexed)
    try:
        selected_A1_idx = [int(x.strip()) - 1 for x in selected_A1.split(',') if x.strip().isdigit()]
        
        # Check if indices are valid
        invalid_indices = [i + 1 for i in selected_A1_idx if i < 0 or i >= len(all_questions)]
        if invalid_indices:
            print(f"Invalid question numbers: {', '.join(map(str, invalid_indices))}")
            input("Press Enter to exit...")
            return
            
        # Get selected questions for Part A - Q1
        selected_A1_questions = [all_questions[i] for i in selected_A1_idx]
        total_marks_A1 = sum(q["marks"] for q in selected_A1_questions)
        
        print(f"\nSelected {len(selected_A1_questions)} questions for Q1, total marks: {total_marks_A1}")
        if abs(total_marks_A1 - 25) > 5:
            print(f"WARNING: Total marks for Q1 ({total_marks_A1}) is not close to 25.")
            proceed = input("Do you want to continue anyway? (y/n): ")
            if proceed.lower() != 'y':
                return
    except Exception as e:
        print(f"Error processing question selection: {e}")
        input("Press Enter to exit...")
        return
    
    # Remove selected questions from the available pool for Q2
    available_for_A2 = [q for i, q in enumerate(all_questions) if i not in selected_A1_idx]
    
    # Display remaining questions for Q2 selection
    print("\nRemaining Questions for Q2:")
    for idx, q in enumerate(available_for_A2, start=1):
        print(f"{idx}. {q['question']} [{q['marks']} marks]")
    
    # Let the user choose questions for PART A (Q2)
    print("\n" + "=" * 50)
    print("PART A - Question 2 Selection")
    print("=" * 50)
    selected_A2 = input("Enter question numbers for Q2 (comma separated, total ~25 marks): ")
    
    # Convert input to indices (0-indexed in the available_for_A2 list)
    try:
        selected_A2_idx = [int(x.strip()) - 1 for x in selected_A2.split(',') if x.strip().isdigit()]
        
        # Check if indices are valid
        invalid_indices = [i + 1 for i in selected_A2_idx if i < 0 or i >= len(available_for_A2)]
        if invalid_indices:
            print(f"Invalid question numbers: {', '.join(map(str, invalid_indices))}")
            input("Press Enter to exit...")
            return
            
        # Get selected questions for Part A - Q2
        selected_A2_questions = [available_for_A2[i] for i in selected_A2_idx]
        total_marks_A2 = sum(q["marks"] for q in selected_A2_questions)
        
        print(f"\nSelected {len(selected_A2_questions)} questions for Q2, total marks: {total_marks_A2}")
        if abs(total_marks_A2 - 25) > 5:
            print(f"WARNING: Total marks for Q2 ({total_marks_A2}) is not close to 25.")
            proceed = input("Do you want to continue anyway? (y/n): ")
            if proceed.lower() != 'y':
                return
    except Exception as e:
        print(f"Error processing question selection: {e}")
        input("Press Enter to exit...")
        return
    
    # Remove selected questions from the available pool for PART B
    used_indices = selected_A1_idx + [selected_A1_idx[-1] + i + 1 for i in selected_A2_idx]
    available_for_B = [q for i, q in enumerate(all_questions) if i not in used_indices]
    
    # Display remaining questions for PART B
    print("\nRemaining Questions for PART B:")
    for idx, q in enumerate(available_for_B, start=1):
        print(f"{idx}. {q['question']} [{q['marks']} marks]")
    
    # Let the user choose questions for PART B (Q3)
    print("\n" + "=" * 50)
    print("PART B - Question 3 Selection")
    print("=" * 50)
    selected_B1 = input("Enter question numbers for Q3 (comma separated, total ~25 marks): ")
    
    # Convert input to indices (0-indexed in the available_for_B list)
    try:
        selected_B1_idx = [int(x.strip()) - 1 for x in selected_B1.split(',') if x.strip().isdigit()]
        
        # Check if indices are valid
        invalid_indices = [i + 1 for i in selected_B1_idx if i < 0 or i >= len(available_for_B)]
        if invalid_indices:
            print(f"Invalid question numbers: {', '.join(map(str, invalid_indices))}")
            input("Press Enter to exit...")
            return
            
        # Get selected questions for Part B - Q3
        selected_B1_questions = [available_for_B[i] for i in selected_B1_idx]
        total_marks_B1 = sum(q["marks"] for q in selected_B1_questions)
        
        print(f"\nSelected {len(selected_B1_questions)} questions for Q3, total marks: {total_marks_B1}")
        if abs(total_marks_B1 - 25) > 5:
            print(f"WARNING: Total marks for Q3 ({total_marks_B1}) is not close to 25.")
            proceed = input("Do you want to continue anyway? (y/n): ")
            if proceed.lower() != 'y':
                return
    except Exception as e:
        print(f"Error processing question selection: {e}")
        input("Press Enter to exit...")
        return
    
    # Remove selected questions from the available pool for Q4
    available_for_B2 = [q for i, q in enumerate(available_for_B) if i not in selected_B1_idx]
    
    # Display remaining questions for Q4 selection
    print("\nRemaining Questions for Q4:")
    for idx, q in enumerate(available_for_B2, start=1):
        print(f"{idx}. {q['question']} [{q['marks']} marks]")
    
    # Let the user choose questions for PART B (Q4)
    print("\n" + "=" * 50)
    print("PART B - Question 4 Selection")
    print("=" * 50)
    selected_B2 = input("Enter question numbers for Q4 (comma separated, total ~25 marks): ")
    
    # Convert input to indices (0-indexed in the available_for_B2 list)
    try:
        selected_B2_idx = [int(x.strip()) - 1 for x in selected_B2.split(',') if x.strip().isdigit()]
        
        # Check if indices are valid
        invalid_indices = [i + 1 for i in selected_B2_idx if i < 0 or i >= len(available_for_B2)]
        if invalid_indices:
            print(f"Invalid question numbers: {', '.join(map(str, invalid_indices))}")
            input("Press Enter to exit...")
            return
            
        # Get selected questions for Part B - Q4
        selected_B2_questions = [available_for_B2[i] for i in selected_B2_idx]
        total_marks_B2 = sum(q["marks"] for q in selected_B2_questions)
        
        print(f"\nSelected {len(selected_B2_questions)} questions for Q4, total marks: {total_marks_B2}")
        if abs(total_marks_B2 - 25) > 5:
            print(f"WARNING: Total marks for Q4 ({total_marks_B2}) is not close to 25.")
            proceed = input("Do you want to continue anyway? (y/n): ")
            if proceed.lower() != 'y':
                return
    except Exception as e:
        print(f"Error processing question selection: {e}")
        input("Press Enter to exit...")
        return
    
    # Get the template file
    template_path = select_file("Select a template file:")
    if not template_path:
        print("No template file selected. Exiting.")
        return
    
    # Load the template DOCX file
    try:
        template_doc = Document(template_path)
        if not template_doc.tables:
            print("Error: Template file does not contain a table.")
            input("Press Enter to exit...")
            return
            
        table = template_doc.tables[0]
    except Exception as e:
        print(f"Error loading template: {e}")
        input("Press Enter to exit...")
        return
    
    # Remove existing rows except the header row
    try:
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[1]._tr)
    except Exception as e:
        print(f"Error preparing template table: {e}")
        input("Press Enter to exit...")
        return
    
    # Function to add a merged header row to the table
    def add_section_header(table, header_text):
        row = table.add_row()
        merged_cell = row.cells[0]
        # Merge all cells in the row to create a section header
        for cell in row.cells[1:]:
            merged_cell = merged_cell.merge(cell)
        merged_cell.text = header_text
    
    try:
        # Add PART A section header
        add_section_header(table, "PART A")
        
        # Add Question 1 with subparts
        q1_label = "1"
        for idx, q in enumerate(selected_A1_questions):
            row_cells = table.add_row().cells
            if idx == 0:
                # First row of Q1
                row_cells[0].text = q1_label
            else:
                # Subparts of Q1
                row_cells[0].text = f"{q1_label}{chr(97+idx)}"
            
            row_cells[1].text = q["question"]
            row_cells[2].text = str(q["marks"])
            row_cells[3].text = q["co"]
            row_cells[4].text = q["rbt"]
        
        # Add Question 2 with subparts
        q2_label = "2"
        for idx, q in enumerate(selected_A2_questions):
            row_cells = table.add_row().cells
            if idx == 0:
                # First row of Q2
                row_cells[0].text = q2_label
            else:
                # Subparts of Q2
                row_cells[0].text = f"{q2_label}{chr(97+idx)}"
            
            row_cells[1].text = q["question"]
            row_cells[2].text = str(q["marks"])
            row_cells[3].text = q["co"]
            row_cells[4].text = q["rbt"]
        
        # Add PART B section header
        add_section_header(table, "PART B")
        
        # Add Question 3 with subparts
        q3_label = "3"
        for idx, q in enumerate(selected_B1_questions):
            row_cells = table.add_row().cells
            if idx == 0:
                # First row of Q3
                row_cells[0].text = q3_label
            else:
                # Subparts of Q3
                row_cells[0].text = f"{q3_label}{chr(97+idx)}"
            
            row_cells[1].text = q["question"]
            row_cells[2].text = str(q["marks"])
            row_cells[3].text = q["co"]
            row_cells[4].text = q["rbt"]
        
        # Add Question 4 with subparts
        q4_label = "4"
        for idx, q in enumerate(selected_B2_questions):
            row_cells = table.add_row().cells
            if idx == 0:
                # First row of Q4
                row_cells[0].text = q4_label
            else:
                # Subparts of Q4
                row_cells[0].text = f"{q4_label}{chr(97+idx)}"
            
            row_cells[1].text = q["question"]
            row_cells[2].text = str(q["marks"])
            row_cells[3].text = q["co"]
            row_cells[4].text = q["rbt"]
        
    except Exception as e:
        print(f"Error creating question paper: {e}")
        input("Press Enter to exit...")
        return
    
    # Get output filename
    output_filename = input("\nEnter the output filename (e.g., 'generated_question_paper.docx'): ")
    if not output_filename.endswith('.docx'):
        output_filename += '.docx'
    
    # Save the generated question paper
    try:
        template_doc.save(output_filename)
        print(f"\nQuestion paper saved as: {os.path.abspath(output_filename)}")
    except Exception as e:
        print(f"Error saving question paper: {e}")
        input("Press Enter to exit...")
        return
    
    print("\nQuestion paper generation complete!")
    input("Press Enter to exit...")

if __name__ == "__main__":
    main()